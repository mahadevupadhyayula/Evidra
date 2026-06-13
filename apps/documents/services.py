from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from django.conf import settings
from django.core.files.base import ContentFile
from django.db import transaction
from django.http import Http404
from docx import Document as DocxDocument
from pypdf import PdfReader

from apps.documents.models import Document, DocumentParsingStatus, DocumentType
from apps.documents.storage import build_resume_storage_key, private_resume_storage
from apps.sprints.models import InterviewSprint
from apps.sprints.services import SprintWorkflowService

PDF_MIME_TYPE = "application/pdf"
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
ALLOWED_RESUME_MIME_TYPES = {PDF_MIME_TYPE, DOCX_MIME_TYPE}
ALLOWED_RESUME_EXTENSIONS = {".pdf", ".docx"}
MIN_RESUME_TEXT_LENGTH = 100
MAX_RESUME_TEXT_LENGTH = 200_000


class InvalidResumeFile(ValueError):
    """Raised when an uploaded resume fails deterministic validation."""


class ResumeParsingFailed(ValueError):
    """Raised when a resume file cannot produce reviewable text without OCR."""


class ResumeConfirmationError(ValueError):
    """Raised when a resume cannot be confirmed safely."""


@dataclass(frozen=True)
class ResumeFileMetadata:
    original_filename: str
    extension: str
    mime_type: str
    file_size: int


class ResumeParserService:
    @staticmethod
    def validate_file(uploaded_file) -> ResumeFileMetadata:
        original_filename = Path(uploaded_file.name or "").name
        extension = Path(original_filename).suffix.lower()
        mime_type = getattr(uploaded_file, "content_type", "") or ""
        file_size = int(getattr(uploaded_file, "size", 0) or 0)

        if not original_filename:
            raise InvalidResumeFile("Choose a PDF or DOCX resume to upload.")
        if extension not in ALLOWED_RESUME_EXTENSIONS:
            raise InvalidResumeFile("Resume must be a PDF or DOCX file.")
        if mime_type not in ALLOWED_RESUME_MIME_TYPES:
            raise InvalidResumeFile("Resume file type is not supported.")
        if file_size <= 0:
            raise InvalidResumeFile("Resume file is empty.")
        if file_size > settings.RESUME_MAX_UPLOAD_BYTES:
            raise InvalidResumeFile("Resume file is larger than the allowed size.")

        ResumeParserService._validate_file_signature(uploaded_file, extension)
        return ResumeFileMetadata(
            original_filename=original_filename,
            extension=extension,
            mime_type=mime_type,
            file_size=file_size,
        )

    @staticmethod
    def _validate_file_signature(uploaded_file, extension: str) -> None:
        position = uploaded_file.tell() if hasattr(uploaded_file, "tell") else None
        try:
            uploaded_file.seek(0)
            if extension == ".pdf":
                if uploaded_file.read(4) != b"%PDF":
                    raise InvalidResumeFile("Uploaded PDF does not look like a PDF file.")
            elif extension == ".docx":
                try:
                    with ZipFile(uploaded_file) as archive:
                        names = set(archive.namelist())
                except BadZipFile as exc:
                    raise InvalidResumeFile(
                        "Uploaded DOCX does not look like a DOCX file."
                    ) from exc
                if "word/document.xml" not in names:
                    raise InvalidResumeFile("Uploaded DOCX is missing document content.")
        finally:
            if position is not None:
                uploaded_file.seek(position)

    @staticmethod
    def extract_pdf_text(file_or_path) -> str:
        try:
            reader = PdfReader(file_or_path)
            raw_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:  # noqa: BLE001
            raise ResumeParsingFailed("We could not extract text from this PDF.") from exc
        return ResumeParserService._require_reviewable_text(raw_text)

    @staticmethod
    def extract_docx_text(file_or_path) -> str:
        try:
            doc = DocxDocument(file_or_path)
            parts = [paragraph.text for paragraph in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            raw_text = "\n".join(parts)
        except Exception as exc:  # noqa: BLE001
            raise ResumeParsingFailed("We could not extract text from this DOCX.") from exc
        return ResumeParserService._require_reviewable_text(raw_text)

    @staticmethod
    def clean_text(raw_text: str) -> str:
        text = raw_text.replace("\x00", "")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def detect_sections(cleaned_text: str) -> dict[str, bool]:
        lower_text = cleaned_text.lower()
        return {
            "experience": any(label in lower_text for label in ["experience", "employment"]),
            "education": "education" in lower_text,
            "skills": "skills" in lower_text,
            "projects": "projects" in lower_text,
        }

    @staticmethod
    def _require_reviewable_text(raw_text: str) -> str:
        cleaned_text = ResumeParserService.clean_text(raw_text)
        if len(cleaned_text) < MIN_RESUME_TEXT_LENGTH:
            raise ResumeParsingFailed(
                "We could not extract enough text from this file. Please paste your resume text."
            )
        return cleaned_text


class ResumeDocumentService:
    @staticmethod
    def get_owned_document(user, document_id) -> Document:
        if not user.is_authenticated:
            raise Http404("Resume not found.")
        try:
            return Document.objects.get(
                pk=document_id,
                user=user,
                document_type=DocumentType.RESUME,
            )
        except Document.DoesNotExist as exc:
            raise Http404("Resume not found.") from exc

    @staticmethod
    def get_active_resume(user) -> Document | None:
        if not user.is_authenticated:
            return None
        return Document.objects.filter(
            user=user,
            document_type=DocumentType.RESUME,
            is_active=True,
        ).first()

    @staticmethod
    def create_from_upload(*, user, uploaded_file) -> Document:
        metadata = ResumeParserService.validate_file(uploaded_file)
        storage = private_resume_storage()
        storage_key = build_resume_storage_key(
            user_id=user.id,
            original_filename=metadata.original_filename,
        )
        uploaded_file.seek(0)
        saved_key = storage.save(storage_key, ContentFile(uploaded_file.read()))

        document = Document.objects.create(
            user=user,
            document_type=DocumentType.RESUME,
            original_filename=metadata.original_filename,
            storage_key=saved_key,
            mime_type=metadata.mime_type,
            file_size=metadata.file_size,
            parsing_status=DocumentParsingStatus.PARSING,
        )

        try:
            path = storage.path(saved_key)
            if metadata.extension == ".pdf":
                raw_text = ResumeParserService.extract_pdf_text(path)
            else:
                raw_text = ResumeParserService.extract_docx_text(path)
            cleaned_text = ResumeParserService.clean_text(raw_text)
        except ResumeParsingFailed as exc:
            document.parsing_status = DocumentParsingStatus.PARSING_FAILED
            document.parsing_error = str(exc)
            document.save(update_fields=["parsing_status", "parsing_error", "updated_at"])
            return document

        document.raw_text = raw_text
        document.cleaned_text = cleaned_text
        document.parsing_status = DocumentParsingStatus.READY_FOR_REVIEW
        document.parsing_error = ""
        document.save(
            update_fields=[
                "raw_text",
                "cleaned_text",
                "parsing_status",
                "parsing_error",
                "updated_at",
            ]
        )
        return document

    @staticmethod
    def create_from_paste(*, user, text: str) -> Document:
        cleaned_text = ResumeDocumentService.normalize_resume_text(text)
        return Document.objects.create(
            user=user,
            document_type=DocumentType.RESUME,
            original_filename="Pasted resume text",
            raw_text=cleaned_text,
            cleaned_text=cleaned_text,
            mime_type="text/plain",
            file_size=len(cleaned_text.encode()),
            parsing_status=DocumentParsingStatus.READY_FOR_REVIEW,
        )

    @staticmethod
    def update_review_text(*, user, document_id, cleaned_text: str) -> Document:
        document = ResumeDocumentService.get_owned_document(user, document_id)
        if document.parsing_status != DocumentParsingStatus.READY_FOR_REVIEW:
            raise ResumeConfirmationError(
                "Confirmed resumes cannot be edited. Replace the resume to make changes."
            )
        document.cleaned_text = ResumeDocumentService.normalize_resume_text(cleaned_text)
        document.parsing_status = DocumentParsingStatus.READY_FOR_REVIEW
        document.save(update_fields=["cleaned_text", "parsing_status", "updated_at"])
        return document

    @staticmethod
    def confirm_resume(*, user, sprint: InterviewSprint, document_id) -> Document:
        with transaction.atomic():
            document = (
                Document.objects.select_for_update()
                .filter(pk=document_id, user=user, document_type=DocumentType.RESUME)
                .first()
            )
            if document is None:
                raise Http404("Resume not found.")
            if document.parsing_status not in {
                DocumentParsingStatus.READY_FOR_REVIEW,
                DocumentParsingStatus.CONFIRMED,
            }:
                raise ResumeConfirmationError("Review resume text before confirming it.")
            if len(document.cleaned_text.strip()) < MIN_RESUME_TEXT_LENGTH:
                raise ResumeConfirmationError("Resume text is too short to confirm.")

            Document.objects.select_for_update().filter(
                user=user,
                document_type=DocumentType.RESUME,
                is_active=True,
            ).exclude(pk=document.pk).update(
                is_active=False,
                parsing_status=DocumentParsingStatus.REPLACED,
            )
            document.is_active = True
            document.parsing_status = DocumentParsingStatus.CONFIRMED
            document.parsing_error = ""
            document.save(
                update_fields=["is_active", "parsing_status", "parsing_error", "updated_at"]
            )

            SprintWorkflowService.mark_resume_ready(user=user, sprint=sprint, document=document)
            return document

    @staticmethod
    def normalize_resume_text(text: str) -> str:
        cleaned_text = ResumeParserService.clean_text(text)
        if len(cleaned_text) < MIN_RESUME_TEXT_LENGTH:
            raise ResumeConfirmationError("Resume text must be at least 100 characters.")
        if len(cleaned_text) > MAX_RESUME_TEXT_LENGTH:
            raise ResumeConfirmationError("Resume text is too long.")
        return cleaned_text
